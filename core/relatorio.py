import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def gerar_relatorio_pt(titulo, conteudo):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    header = doc.add_heading("", level=0)
    run_h = header.add_run("J.A.R.V.I.S — STARK INDUSTRIES")
    run_h.font.size = Pt(13)
    run_h.font.color.rgb = RGBColor(0, 100, 150)
    run_h.font.bold = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    titulo_par = doc.add_heading("", level=1)
    run_t = titulo_par.add_run(titulo.upper())
    run_t.font.size = Pt(20)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0, 70, 120)
    titulo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data_par = doc.add_paragraph()
    data_run = data_par.add_run(
        f"Gerado em: {datetime.datetime.now().strftime('%d/%m/%Y às %H:%M')}"
    )
    data_run.font.size = Pt(9)
    data_run.font.color.rgb = RGBColor(120, 120, 120)
    data_run.font.italic = True
    data_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sep = doc.add_paragraph("─" * 60)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.runs[0].font.color.rgb = RGBColor(0, 150, 200)
    sep.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    for linha in conteudo.split("\n"):
        linha = linha.strip()
        if not linha:
            doc.add_paragraph()
            continue

        if linha.startswith(("•", "-", "*")):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(linha.lstrip("•-* "))
            run.font.size = Pt(11)


        elif linha.isupper() and len(linha) > 3:
            p = doc.add_heading("", level=2)
            run = p.add_run(linha)
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 100, 150)


        else:
            p = doc.add_paragraph()
            run = p.add_run(linha)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(30, 30, 30)

    doc.add_paragraph()

    rodape = doc.add_paragraph("─" * 60)
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rodape.runs[0].font.color.rgb = RGBColor(0, 150, 200)
    rodape.runs[0].font.size = Pt(10)

    fim = doc.add_paragraph()
    fim_run = fim.add_run("Documento gerado automaticamente pelo sistema J.A.R.V.I.S  ·  CONFIDENCIAL")
    fim_run.font.size = Pt(8)
    fim_run.font.color.rgb = RGBColor(150, 150, 150)
    fim_run.font.italic = True
    fim.alignment = WD_ALIGN_PARAGRAPH.CENTER

    nome_ficheiro = f"{titulo.replace(' ', '_')}.docx"
    doc.save(nome_ficheiro)
    os.startfile(nome_ficheiro)

    return nome_ficheiro

